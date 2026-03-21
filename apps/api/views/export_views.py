from django.http import HttpResponse, Http404
from django.template.loader import render_to_string
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
from django.views import View
from django.template.loader import get_template
import logging
# WeasyPrint 延迟导入，避免启动时依赖问题
# from weasyprint import HTML
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from apps.models.itinerary import Itinerary


class ItineraryPDFExportView(View):
    """
    Exports an itinerary as a PDF using WeasyPrint.
    URL pattern should pass itinerary_id as a path parameter.
    """
    logger = logging.getLogger(__name__)

    @method_decorator(csrf_exempt)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, itinerary_id, *args, **kwargs):
        # 延迟导入 WeasyPrint
        try:
            from weasyprint import HTML
        except OSError as e:
            self.logger.error(f"WeasyPrint 库未正确安装: {e}")
            raise Http404("PDF导出功能暂时不可用，请联系管理员")
        
        try:
            # Fetch itinerary with related data
            itinerary = Itinerary.objects.get(itinerary_id=itinerary_id)
        except Itinerary.DoesNotExist:
            self.logger.exception("Itinerary not found: id=%s", itinerary_id)
            raise Http404(f"Itinerary not found: {itinerary_id}")
        except Exception:
            self.logger.exception("Error loading itinerary: id=%s", itinerary_id)
            raise Http404("Error loading itinerary.")

        # Prepare context using the same data as preview_itinerary view
        try:
            destinations = list(itinerary.destinations.all())
            schedules = itinerary.dailyschedule_set.all().order_by('schedule_date', 'start_time')

            grouped_schedules = {}
            for ds in schedules:
                key = ds.schedule_date.isoformat() if ds.schedule_date else ''
                grouped_schedules.setdefault(key, []).append(ds)

            traveler_stats_list = list(itinerary.traveler_stats.all())

            context = {
                'itinerary': itinerary,
                'destinations': destinations,
                'grouped_schedules': grouped_schedules,
                'traveler_stats': traveler_stats_list[0] if traveler_stats_list else None,
            }

            # Render template to string
            rendered = render_to_string('admin/preview_itinerary.html', context=context, request=request)

            # Generate PDF using WeasyPrint
            pdf_bytes = HTML(string=rendered, base_url=request.build_absolute_uri('/')).write_pdf()

            filename = f"itinerary_{itinerary_id}.pdf"
            response = HttpResponse(pdf_bytes, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        except Exception:
            self.logger.exception("Failed to export itinerary PDF: id=%s", itinerary_id)
            raise Http404("Failed to export itinerary PDF.")


class ItineraryWordExportView(View):
    """
    使用 Word 模板导出行程
    模板文件: templates/export/itinerary_template.docx
    """
    logger = logging.getLogger(__name__)

    @method_decorator(csrf_exempt)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, itinerary_id, *args, **kwargs):
        try:
            itinerary = Itinerary.objects.get(itinerary_id=itinerary_id)
        except Itinerary.DoesNotExist:
            raise Http404(f"Itinerary not found: {itinerary_id}")

        # 加载模板
        try:
            doc = Document('templates/export/itinerary_template.docx')
        except Exception:
            raise Http404("Word template not found")

        # 设置默认字体为微软雅黑
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        # 设置中文字体回退
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # 准备数据
        try:
            traveler_stats = itinerary.traveler_stats.first()
            total_travelers = 0
            adult_count = 0
            child_count = 0
            infant_count = 0
            if traveler_stats:
                total_travelers = traveler_stats.adult_count + traveler_stats.child_count + traveler_stats.infant_count + traveler_stats.senior_count
                adult_count = traveler_stats.adult_count
                child_count = traveler_stats.child_count
                infant_count = traveler_stats.infant_count
        except Exception:
            pass

        # 目的地
        destinations = itinerary.destinations.all()
        destinations_text = ''
        for dest in destinations:
            destinations_text += f"{dest.city_name} - {dest.arrival_date} 至 {dest.departure_date} ({dest.nights} 晚)\n"

        # 每日行程
        schedules = itinerary.dailyschedule_set.all().order_by('schedule_date', 'start_time')
        schedules_text = ''
        current_date = None
        for schedule in schedules:
            date_str = str(schedule.schedule_date) if schedule.schedule_date else ''
            if date_str != current_date:
                current_date = date_str
                schedules_text += f"\n第 {schedule.day_number} 天 - {date_str}\n"
                schedules_text += "=" * 30 + "\n"
            
            time_str = f"{schedule.start_time} - {schedule.end_time}" if schedule.start_time and schedule.end_time else ''
            if time_str:
                schedules_text += f"  {time_str} "
            schedules_text += f"{schedule.activity_title or ''}\n"
            if schedule.activity_description:
                schedules_text += f"    {schedule.activity_description}\n"
            if schedule.destination_id:
                schedules_text += f"    地点: {schedule.destination_id.city_name}\n"

        # 替换占位符
        replacements = {
            '{{itinerary_name}}': itinerary.itinerary_name or '',
            '{{itinerary_id}}': itinerary.itinerary_id or '',
            '{{total_days}}': str(itinerary.total_days) if itinerary.total_days else '',
            '{{start_date}}': str(itinerary.start_date) if itinerary.start_date else '',
            '{{end_date}}': str(itinerary.end_date) if itinerary.end_date else '',
            '{{departure_city}}': itinerary.departure_city or '',
            '{{return_city}}': itinerary.return_city or '',
            '{{contact_person}}': itinerary.contact_person or '',
            '{{contact_phone}}': itinerary.contact_phone or '',
            '{{total_budget}}': str(itinerary.total_budget) if itinerary.total_budget else '',
            '{{current_status}}': str(itinerary.current_status) if itinerary.current_status else '',
            '{{total_travelers}}': str(total_travelers),
            '{{adult_count}}': str(adult_count),
            '{{child_count}}': str(child_count),
            '{{infant_count}}': str(infant_count),
            '{{destinations}}': destinations_text.strip(),
            '{{daily_schedules}}': schedules_text.strip(),
            '{{description}}': itinerary.description or '',
            '{{itinerary_quote}}': itinerary.itinerary_quote or '',
            '{{fee_included}}': '',
            '{{fee_excluded}}': '',
            '{{special_notes}}': '',
        }

        # 遍历所有段落，替换文本
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    # 保留原始格式，只替换文本
                    inline = paragraph.runs
                    for run in inline:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

        # 遍历所有表格，替换文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                inline = paragraph.runs
                                for run in inline:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)

        # 保存到响应
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        filename = f"itinerary_{itinerary_id}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
